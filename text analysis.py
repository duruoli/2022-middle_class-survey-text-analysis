# -*- coding: utf-8 -*-
"""
Created on Tue Apr 13 14:00:02 2021

@author: ldr
"""



from win32com import client as wc
import os #用于获取目标文件所在路径
import docx
import sys
import pandas as pd
import numpy as np
import jieba

##1 提取直接相关的问卷问题（“中产”）
os.chdir(r"D:\A李杜若\留学\科研\范老师中产项目\practise\A") #创建新文件的路径
filenames=os.listdir(r"D:\A李杜若\留学\科研\范老师中产项目\practise")
filenames1=filenames[0:5]

k=0
for fnm in filenames1:
    log_d="D:\\A李杜若\\留学\\科研\\范老师中产项目\\practise"
    logFiles=os.listdir(log_d+'\\'+fnm)
    for file in logFiles:
        q=docx.Document(log_d+'\\'+fnm+'\\'+file)
        name=os.path.basename(file).replace('.docx','')
        fn = open("%s.txt"%name,'w')     #直接打开一个文件，如果文件不存在则创建文件
        k=0
        for i in range(len(q.paragraphs)):
            if '中产'in q.paragraphs[i].text and ('概念'in q.paragraphs[i].text or '理解'in q.paragraphs[i].text) :
                k=1
            if '下层'in q.paragraphs[i].text and '关系' in q.paragraphs[i].text and '？' in q.paragraphs[i].text:
                k=2
            if k==1:
                fn.write(q.paragraphs[i].text+'\n')
            elif k==2:
                fn.write(q.paragraphs[i].text+'\n'+q.paragraphs[i+1].text)
                fn.close()
                break

#--4.20--
## 2 整理出dataframe
#手动整理数据/小丑竟是我自己.../                
#提取基本信息
df0=pd.DataFrame(columns = ('编号','调查日期','性别','年龄','工作性质','移民与否'))
df0.set_index('编号', inplace = True)

                
filenames=os.listdir(r"D:\A李杜若\留学\科研\范老师中产项目\practise\A")


for f in filenames:
      q=f.split('-')
      q=q[1:7]
      series=pd.Series({'调查日期':q[1],'性别':q[2],'年龄':q[3],'工作性质':q[4],'移民与否':q[5]}, name=q[0])
      df0=df0.append(series)

#“是否知道中产”变量

df0['是否知道中产？']=[1 for i in range(50)]
list=["A6", "A20", "B6", "B15", "B26"]
for i in list:
    df0.loc[i,['是否知道中产？']]=0

##--4.21--
#处理问卷，统一格式

#删去多余的空行：
for f in filenames:
    with open("%s"%f, "r") as t:
         lines = t.readlines()
    with open("%s"%f,"w") as t_w:
        for line in lines:
            if line == "\n":
                continue
            t_w.write(line)    
 
yl=[]         
nl=[]
for f in filenames:
     with open("%s"%f, "r") as t:
         lines = t.readlines()
         f.replace('.txt','')
     if len(lines)!=18:
         nl.append(f)
     else:
         yl.append(f)
#问了完整的问题的只有五个！！！！！！哭了，其实连五个都没有
#开始补全大业。。。

wid=['理解','属于','算','发展','作用','独特','中产阶层焦虑','有这种','造成','上层','下层']
for f in filenames:
     with open("%s"%f, "r") as t:
         lines = t.readlines()
     with open("新%s"%f,"w") as t_w:
         for index in wid:
             flag=0
             for i in range(0,len(lines)-1,2):
                 if index in lines[i]:
                     t_w.write('Q:'+index+'\n')
                     t_w.write(lines[i+1])
                     flag=1
                     break
             if flag==0:
                 t_w.write('Q:'+index+'\n')
                 t_w.write('A:\n')
         
                     
###一一校正修改：费老鼻子劲儿
os.chdir(r"D:\A李杜若\留学\科研\范老师中产项目\practise\B")
files=os.listdir()
fo=files            
for i in range(len(files)):
      q=files[i].split('-')
      os.rename(files[i],q[1]+'.txt') ##更改文件名为Ai、Bi

filenew=os.listdir() 

df1=df0    
df2=pd.DataFrame(columns=('中产阶级的理解','自我定位','未来发展','对社会的作用','独特之处','知道中产阶级焦虑吗','自身有这种焦虑？','导致焦虑的原因','中产和上层阶级的关系','中产和下层阶级的关系') )   
for f in filenew:
    q=f.replace('.txt','')
    with open("%s"%f, "r") as t:
         lines = t.readlines()
         k=0
         q1=lines
    for line in lines:
        line=line.replace('A:','')
        line=line.replace('B:','')
        line=line.replace('A：','')
        line=line.replace('B：','')
        line=line.replace('B :','')
        line=line.replace('·','')
        line=line.replace('答:','')
        line=line.replace('答：','')
        line=line.replace('-','')
        q1[k]=line
        k=k+1
    dfn=pd.DataFrame({'中产阶级的理解':q1[1]+q1[3],'自我定位':q1[5],'未来发展':q1[7],'对社会的作用':q1[9],'独特之处':q1[11],'知道中产阶级焦虑吗':q1[13],'自身有这种焦虑？':q1[15],'导致焦虑的原因':q1[17],'中产和上层阶级的关系':q1[19],'中产和下层阶级的关系':q1[21]},index=[q])
    df2=pd.concat([df2,dfn],axis=0)

df2_0=df2 #备份df2：采访中的文字数据
df1=df1.join(df2)
df1_0=df1 #备份df1：最终的数据表



##导出表格
df1.to_csv('中产数据集.csv',encoding='gbk')


###文本分析
os.chdir(r"D:\A李杜若\留学\科研\范老师中产项目\practise\B")
df1=pd.read_csv('中产数据集.csv',encoding="GBK")
df1.set_index('编号', inplace = True)
type(df1.iloc[0,7])
s=jieba.cut(df1.iloc[0,7])


os.chdir(r'D:\\A李杜若\\留学\\科研\\范老师中产项目\\stopwords-master')
with open('cn_stopwords.txt', mode="r",encoding='utf-8') as t:
    lines = t.readlines() #导入停用词
segments=[]
jieba.load_userdict('cn_stopwords.txt')
segs=jieba.lcut(df1.iloc[0,6],cut_all=False)
print(' '.join(segs))
for seg in segs:
    segments.append(seg)

###one-hot编码：
##0-1
gender=list(df1['性别'])
for i in range(50):
    if gender[i]=='男':
        gender[i]=1
    else:
        gender[i]=0
    
imm=list(df1['移民与否'])
for i in range(50):
    if imm[i]=='移民':
        imm[i]=1
    else:
        imm[i]=0

##标称型数据，没有顺序和数值计算的意义，只作区分
job=df1['工作性质']
job_dummies=pd.get_dummies(job)#构造变量，从df中抽取出来

###训练和测试
from sklearn.model_selection import train_test_split
from sklearn.tree import DecisionTreeClassifier
from sklearn.model_selection import cross_val_score


####1\是否知道中产
know1=list(df1['是否知道中产？'])
know1_y=np.array(know1).reshape(-1,1)

gender_x=np.array(gender).reshape(-1,1)
imm_x=np.array(imm).reshape(-1,1)
job_x=job_dummies
age=list(df1['年龄'])
age_x=np.array(age).reshape(-1,1)

X=np.concatenate([gender_x,age_x,imm_x,job_x],axis=1)

from sklearn.model_selection import train_test_split
from sklearn.tree import DecisionTreeClassifier
from sklearn.model_selection import cross_val_score

x_train, x_test, y_train, y_test = train_test_split(X, know1_y, test_size=0.3, random_state=2)#分割训练集测试集
#随机森林
tree = DecisionTreeClassifier(max_depth=6,random_state=0) #树的深度设置为6
dt_tree=tree.fit(x_train,y_train)  

score = cross_val_score(tree, X, know1_y, cv=10, scoring='accuracy')
np.mean(score)#10-折交叉验证平均得分 0.82
tree.score(x_train,y_train)#训练集得分 1
tree.score(x_test,y_test)#测试集得分 0.867 

y_predict=dt_tree.predict(x_test)
yy=np.column_stack([y_test,y_predict])#然鹅=0的一个预测对的都没有

print("training set score:{:.3f}".format(tree.score(x_train,y_train)))
print("test set score:{:.3f}".format(tree.score(x_test,y_test)))
print("ten cross-validation score:{:.3f}".format(np.mean(score)))
print("Feature importances : \n{}".format(tree.feature_importances_))

#path="D:\\A李杜若\\留学\\科研\\范老师中产项目\\Program of Social Status\\4 杭州-访谈录音文字稿\\" # 文件夹绝对路径
# files=[]
# for file in os.listdir(path):
#     if file.endswith(".doc"): #排除文件夹内的其它干扰文件，只获取".doc"后缀的word文件
#         files.append(path+file) 
# files

# word = wc.Dispatch("Word.Application") # 打开word应用程序
# for file in files:
#     doc = word.Documents.Open(file) #打开word文件
#     doc.SaveAs("{}x".format(file), 12)#另存为后缀为".docx"的文件，其中参数12指docx文件
#     doc.Close() #关闭原来word文件
# word.Quit()
# print("完成！")




# log_d = "D:\\A李杜若\\留学\\科研\\范老师中产项目\\Program of Social Status\\4 杭州-访谈录音文字稿"
# logFiles = os.listdir(log_d)
#uPath = unicode(cPath,'utf-8')#转码防止出现乱码



    
# #2）单独处理4 无题头问卷，分开每个问题
# os.chdir(r"D:\A李杜若\留学\科研\范老师中产项目\practise\B") #创建新文件的路径    
# logFiles=os.listdir(r"D:\A李杜若\留学\科研\范老师中产项目\practise\4 无题头-问卷-填空式基本情况")
# log2=os.listdir(r"D:\A李杜若\留学\科研\范老师中产项目\practise\A")
# i=0
# s=[0 for i in range(8)]
# for f in logFiles:
#       q=f.split('-')
#       s[i]=q[1]
#       i=i+1
# name = os.listdir
# for file in log2:
#     if any(i in file for i in s):
#         k=open(r"D:\A李杜若\留学\科研\范老师中产项目\practise\A\%s"%file,'r')
#         print(k.read()+'-------------------------') #查看单独处理的问题
    
    

    
# log_d="D:\\A李杜若\\留学\\科研\\范老师中产项目\\practise\\4 无题头-问卷-填空式基本情况"

# k=0
# for file in logFiles:
#         q=docx.Document(log_d+'\\'+file)
#         name=os.path.basename(file).replace('.docx','')
#         fn = open("%s.txt"%name,'w')     #直接打开一个文件，如果文件不存在则创建文件
#         k=0
#         for i in range(len(q.paragraphs)):
#             if '中产'in q.paragraphs[i].text and '理解'in q.paragraphs[i].text:
#                 fn.write(q.paragraphs[i].text)
#                 if '下层' in q.paragraphs[i].text:
#                     break
#                 else:
#                     k=1
#             if '下层'in q.paragraphs[i].text and '关系'in q.paragraphs[i].text and '？' in q.paragraphs[i].text and '理解'not in q.paragraphs[i].text:
#                 k=2
#             if k==1:
#                fn.write(q.paragraphs[i].text)
#             elif k==2:
#                 fn.write(q.paragraphs[i].text+'\n'+q.paragraphs[i+1].text)
#                 fn.close()
#                 break
# ##这个问卷的文字格式真是一言难尽啊！           
# lf=os.listdir(r"D:\A李杜若\留学\科研\范老师中产项目\practise\B")
# for file in lf:
#     i=1              
#     with open("%s"%file, "r+") as f:  # 打开文件
#         s = f.read()# 读取文件
#         sp=s.split("-")
#         sp.remove('')
#     with open("%s"%file, "w") as f1:
#         i=1
#         for a in sp:
#             if i==1:
#                 f1.write('问：'+a+'\n') 
#                 i=i-1
#             elif i ==0:
#                 f1.write('答：'+a+'\n') 
#                 i=i+1
            
            
            
       
            
            
            
     # elif '中产阶层'in q.paragraphs[i-1].text and ('问：' in q.paragraphs[i-1].text or 'Q：'in q.paragraphs[i-1].text or 'A：'in q.paragraphs[i-1].text or '问:' in q.paragraphs[i-1].text or 'Q:'in q.paragraphs[i-1].text):
            #     fn.write(q.paragraphs[i].text+'\n')



